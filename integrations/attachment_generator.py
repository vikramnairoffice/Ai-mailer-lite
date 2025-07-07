"""
PDF, DOCX, and image attachment generation for Email Marketing System
"""

import os
from typing import Dict, Any, Optional
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document

def format_phone_number(phone: str, format_type: str = 'display') -> str:
    """Format phone number for display or click-to-call"""
    if not phone:
        return ""
    
    # Remove all non-digit characters
    digits = ''.join(filter(str.isdigit, phone))
    
    if format_type == 'display':
        if len(digits) == 10:
            return f"({digits[:3]}) {digits[3:6]}-{digits[6:]}"
        elif len(digits) == 11 and digits[0] == '1':
            return f"({digits[1:4]}) {digits[4:7]}-{digits[7:]}"
        else:
            return phone
    elif format_type == 'click_to_call':
        return f"tel:+{digits}"
    
    return phone

def generate_attachment(lead: Dict[str, Any], attachment_type: str, 
                       phone_number: Optional[str] = None) -> Optional[str]:
    """Generate attachment based on type"""
    if attachment_type == 'none':
        return None
    
    phone_config = {'phone_number': phone_number} if phone_number else {}
    
    if attachment_type == 'pdf':
        return create_pdf_attachment(lead, phone_config)
    elif attachment_type == 'docx':
        return create_docx_attachment(lead, phone_config)
    elif attachment_type == 'image':
        # For image, first create PDF then could convert (simplified for now)
        return create_pdf_attachment(lead, phone_config)
    else:
        return create_pdf_attachment(lead, phone_config)

def create_pdf_attachment(lead_data: Dict[str, Any], 
                         phone_config: Dict[str, Any]) -> Optional[str]:
    """Create PDF attachment with contact information"""
    try:
        # Create temp directory
        temp_dir = os.path.join(os.getcwd(), 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # Create PDF file path
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_path = os.path.join(temp_dir, f'contact_info_{timestamp}.pdf')
        
        # Create PDF
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title and content
        story.append(Paragraph("Contact Information", styles['Heading1']))
        story.append(Spacer(1, 20))
        
        if phone_config.get('phone_number'):
            phone_display = format_phone_number(phone_config['phone_number'], 'display')
            story.append(Paragraph(f"<b>Phone:</b> {phone_display}", styles['Normal']))
        
        if lead_data.get('email'):
            story.append(Paragraph(f"<b>Email:</b> {lead_data['email']}", styles['Normal']))
        
        # Add personalized message
        story.append(Spacer(1, 20))
        message = f"Dear {lead_data.get('firstname', 'there')}, Thank you for your interest in our services. Please contact us using the information above. Best regards, Your Name"
        story.append(Paragraph(message, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        
        return pdf_path
        
    except Exception:
        return None

def create_docx_attachment(lead_data: Dict[str, Any], 
                          phone_config: Dict[str, Any]) -> Optional[str]:
    """Create DOCX attachment with contact information"""
    try:
        # Create temp directory
        temp_dir = os.path.join(os.getcwd(), 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # Create DOCX file path
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = os.path.join(temp_dir, f'contact_info_{timestamp}.docx')
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Contact Information', 0)
        title.alignment = 1  # Center alignment
        
        # Add contact details
        doc.add_paragraph()
        
        if phone_config.get('phone_number'):
            phone_display = format_phone_number(phone_config['phone_number'], 'display')
            phone_para = doc.add_paragraph()
            phone_para.add_run('Phone: ').bold = True
            phone_para.add_run(phone_display)
        
        if lead_data.get('email'):
            email_para = doc.add_paragraph()
            email_para.add_run('Email: ').bold = True
            email_para.add_run(lead_data['email'])
        
        # Add personalized message
        doc.add_paragraph()
        message = f"Dear {lead_data.get('firstname', 'there')}, Thank you for your interest in our services. Please contact us using the information above. Best regards, Your Name"
        doc.add_paragraph(message)
        
        # Save document
        doc.save(docx_path)
        
        return docx_path
        
    except Exception:
        return None

def cleanup_temp_files() -> None:
    """Clean up temporary attachment files"""
    try:
        temp_dir = os.path.join(os.getcwd(), 'temp')
        if os.path.exists(temp_dir):
            for filename in os.listdir(temp_dir):
                file_path = os.path.join(temp_dir, filename)
                if os.path.isfile(file_path):
                    os.remove(file_path)
    except Exception:
        pass