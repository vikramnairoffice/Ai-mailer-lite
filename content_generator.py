import random
import re
import os
import json
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
import logging
from jinja2 import Template
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont
import io
import base64
import requests
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import google.generativeai as genai
from utils import format_phone_number, logger

class ContentGenerator:
    """Generate various types of email content."""
    
    def __init__(self, api_key: Optional[str] = None, model_name: Optional[str] = None):
        self.logger = logging.getLogger(__name__)
        self.gemini_model = None
        self.model_name = model_name or "gemini-2.5-flash"
        
        # Initialize Gemini with provided API key or environment variable
        if not api_key:
            api_key = os.getenv('GEMINI_API_KEY')
        
        if api_key:
            try:
                genai.configure(api_key=api_key)
                self.gemini_model = genai.GenerativeModel(self.model_name)
                self.logger.info(f"Gemini AI initialized successfully with {self.model_name}")
            except Exception as e:
                self.logger.error(f"Error initializing Gemini AI: {e}")
    
    def update_api_key(self, api_key: str, model_name: Optional[str] = None) -> bool:
        """Update the API key and optionally model, then reinitialize"""
        if not api_key:
            return False
        
        if model_name:
            self.model_name = model_name
        
        try:
            genai.configure(api_key=api_key)
            self.gemini_model = genai.GenerativeModel(self.model_name)
            self.logger.info(f"Gemini AI reinitialized with new API key and {self.model_name}")
            return True
        except Exception as e:
            self.logger.error(f"Error updating Gemini API key: {e}")
            self.gemini_model = None
            return False
    
    def update_model(self, model_name: str) -> bool:
        """Update the model name and reinitialize if API key exists"""
        if not model_name:
            return False
        
        self.model_name = model_name
        
        # If we already have a working model, reinitialize with new model
        if self.gemini_model:
            try:
                self.gemini_model = genai.GenerativeModel(self.model_name)
                self.logger.info(f"Gemini model updated to {self.model_name}")
                return True
            except Exception as e:
                self.logger.error(f"Error updating Gemini model: {e}")
                return False
        
        return True
    
    def generate_short_spintax(self, lead_data: Dict[str, Any], phone_config: Dict[str, Any], 
                             ai_enhance: bool = False) -> str:
        """Generate short spintax content."""
        
        # Base spintax templates
        templates = [
            "Hi {firstname|there}, {I hope this finds you well|Hope you're doing great}. {I wanted to reach out|Just wanted to connect} about {your business|your company|your work}. {Would love to chat|Let's connect} {soon|when you have time}. {Best regards|Thanks}!",
            
            "{Hey|Hi} {firstname|there}! {Quick question|Just wondering} - {are you looking for|do you need} {better solutions|new opportunities|improvements} for {your business|your company}? {Let's talk|Give me a call} {when convenient|at your convenience}. {Thanks|Best}!",
            
            "{Good morning|Good afternoon} {firstname|there}, {I came across|I found} your {company|business} and {was impressed|thought you might be interested} in {what we offer|our solutions}. {Can we schedule|Would you like} a {quick call|brief chat}? {Looking forward to hearing from you|Hope to connect soon}!",
            
            "{firstname|Hi there}, {I specialize in|We help with} {business growth|improving operations|solving problems} for {companies like yours|businesses in your industry}. {Would you be interested|Are you open} to {learning more|a quick conversation}? {Let me know|Please reach out}!",
            
            "{Hello|Hi} {firstname|there}! {I noticed|I saw} {your company|your business} {online|in my research} and {thought you might benefit|believe you could use} from {our services|what we offer}. {Free to chat|Available for a call} {this week|soon}? {Thanks for your time|Best regards}!"
        ]
        
        # Select random template
        template = random.choice(templates)
        
        # AI enhancement
        if ai_enhance and self.gemini_model:
            try:
                prompt = f"""
                Enhance this spintax email template to make it more engaging and personalized:
                
                Original: {template}
                
                Lead information:
                - Name: {lead_data.get('firstname', 'there')}
                - Company: {lead_data.get('company', 'your company')}
                - Industry: {lead_data.get('industry', 'your industry')}
                
                Requirements:
                - Keep the spintax format with {option1|option2|option3}
                - Make it more personalized and engaging
                - Keep it short (under 100 words)
                - Include at least 3 variations for each section
                - Make it sound natural and professional
                """
                
                response = self.gemini_model.generate_content(prompt)
                enhanced_template = response.text.strip()
                
                # Validate that it's still spintax format
                if '{' in enhanced_template and '|' in enhanced_template:
                    template = enhanced_template
                    
            except Exception as e:
                self.logger.error(f"Error enhancing spintax with AI: {e}")
        
        # Process spintax and personalize
        content = self._process_spintax(template, lead_data, phone_config)
        
        return content
    
    def generate_long_spintax(self, lead_data: Dict[str, Any], phone_config: Dict[str, Any], 
                            ai_enhance: bool = False) -> str:
        """Generate long spintax content."""
        
        # Extended spintax templates
        templates = [
            """
            {Hi|Hello|Good morning|Good afternoon} {firstname|there},

            {I hope this email finds you well|Hope you're having a great day|Trust this message reaches you in good spirits}. {I'm reaching out|I wanted to connect|I'm contacting you} because {I came across|I discovered|I found} {your company|your business|{company}} {online|in my research|during my search} and {was impressed by|really liked|found interesting} {what you do|your work|your services}.

            {I specialize in|We focus on|Our company helps with} {helping businesses|supporting companies|assisting organizations} {like yours|in your industry|in the {industry} sector} {grow and succeed|achieve their goals|reach new heights|overcome challenges}. {We've helped|I've assisted|Our team has supported} {hundreds of|many|numerous} {businesses|companies|organizations} {improve their|enhance their|optimize their} {operations|processes|performance|results}.

            {I'd love to|Would love to|I'm interested in} {learn more about|understand better|discuss} {your current|your existing|the} {challenges|situation|needs|goals} and {see how|explore how|determine if} {we can help|I can assist|our solutions might benefit} {you|your business|your team}. {This could be|It might be|There's potential for} {a great fit|an excellent opportunity|a perfect match} {for both of us|for your business|for your goals}.

            {Would you be|Are you|Might you be} {interested in|open to|available for} {a quick call|a brief conversation|a short chat|a 15-minute discussion} {this week|in the coming days|when convenient|at your convenience}? {I can work around|I'm flexible with|Let me know} your schedule.

            {Looking forward to|Hope to|Excited to} {hearing from you|connecting with you|our conversation}!

            {Best regards|Thanks|Best wishes|Warm regards},
            {Your Name|[Your Name]|[Sender Name]}
            """,
            
            """
            {Dear|Hello|Hi} {firstname|there},

            {I hope you don't mind|Hope it's okay|Trust you don't mind} me reaching out {directly|like this|via email}. {I'm|We're} {currently working with|helping|supporting} {several|many|numerous} {businesses|companies} {in your area|in the {industry} industry|similar to yours} and {thought you might|believed you could|felt you would} {benefit from|be interested in|find value in} {what we offer|our services|our solutions}.

            {The reason I'm contacting you|Why I'm reaching out|The purpose of this email} is {simple|straightforward|clear}: {we've developed|I've created|our team has built} {a system|a solution|a process} that {helps|allows|enables} {businesses|companies} {like yours|in your position|facing similar challenges} to {significantly improve|dramatically enhance|substantially increase} their {results|performance|outcomes|success}.

            {Here's what makes us different|What sets us apart|Our unique approach}: {we don't just|we avoid|we never} {talk theory|give generic advice|use cookie-cutter solutions}. {Instead|Rather|Instead of that}, {we focus on|we concentrate on|we prioritize} {practical|real-world|actionable} {solutions|strategies|approaches} that {deliver|produce|generate} {real results|tangible outcomes|measurable improvements} {quickly|fast|in a short time}.

            {I'd be happy to|Would love to|I'm excited to} {share more details|explain further|provide more information} about {how this works|our approach|what we do} and {answer any questions|address any concerns|discuss your specific situation} you might have.

            {Could we|Would it be possible to|Are you available to} {schedule|arrange|set up} {a quick call|a brief phone conversation|a short discussion} {sometime this week|in the next few days|when it's convenient for you}?

            {Thank you for your time|Thanks for reading|I appreciate your attention},
            {Best|Regards|Best regards},
            {[Your Name]|Your Name}
            """
        ]
        
        # Select random template
        template = random.choice(templates)
        
        # AI enhancement
        if ai_enhance and self.gemini_model:
            try:
                prompt = f"""
                Enhance this long-form spintax email template to make it more compelling and personalized:
                
                Original: {template}
                
                Lead information:
                - Name: {lead_data.get('firstname', 'there')}
                - Company: {lead_data.get('company', 'your company')}
                - Industry: {lead_data.get('industry', 'your industry')}
                - Location: {lead_data.get('location', 'your area')}
                
                Requirements:
                - Keep the spintax format with {option1|option2|option3}
                - Make it more personalized and compelling
                - Add industry-specific language when possible
                - Include at least 3-4 variations for each section
                - Keep professional but engaging tone
                - Length should be 200-300 words when processed
                """
                
                response = self.gemini_model.generate_content(prompt)
                enhanced_template = response.text.strip()
                
                # Validate that it's still spintax format
                if '{' in enhanced_template and '|' in enhanced_template:
                    template = enhanced_template
                    
            except Exception as e:
                self.logger.error(f"Error enhancing long spintax with AI: {e}")
        
        # Process spintax and personalize
        content = self._process_spintax(template, lead_data, phone_config)
        
        return content
    
    def generate_html_template(self, lead_data: Dict[str, Any], phone_config: Dict[str, Any], 
                             ai_enhance: bool = False) -> str:
        """Generate HTML email template."""
        
        # Base HTML templates
        templates = [
            """
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    body { font-family: Arial, sans-serif; line-height: 1.6; color: #333; }
                    .container { max-width: 600px; margin: 0 auto; padding: 20px; }
                    .header { background: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 20px; }
                    .content { padding: 20px 0; }
                    .cta { background: #007bff; color: white; padding: 12px 24px; text-decoration: none; border-radius: 4px; display: inline-block; margin: 20px 0; }
                    .footer { border-top: 1px solid #eee; padding-top: 20px; margin-top: 30px; color: #666; font-size: 14px; }
                </style>
            </head>
            <body>
                <div class="container">
                    <div class="header">
                        <h2>Hello {{ firstname }}!</h2>
                    </div>
                    <div class="content">
                        <p>I hope this email finds you well. I'm reaching out because I came across {{ company }} and was impressed by what you do in the {{ industry }} industry.</p>
                        
                        <p>We specialize in helping businesses like yours achieve their goals and overcome challenges. Our proven approach has helped hundreds of companies improve their operations and results.</p>
                        
                        <p>I'd love to learn more about your current situation and see how we might be able to help {{ company }} reach new heights.</p>
                        
                        <div class="cta">Call Us: {{ phone_display }} (Please save this number for future contact)</div>
                        
                        <p>Would you be interested in a quick 15-minute conversation this week? I can work around your schedule.</p>
                        
                        <p>Looking forward to hearing from you!</p>
                    </div>
                    <div class="footer">
                        <p>Best regards,<br>Your Name</p>
                        {% if phone_display %}
                        <p>📞 {{ phone_display }}</p>
                        {% endif %}
                    </div>
                </div>
            </body>
            </html>
            """,
            
            """
            <!DOCTYPE html>
            <html>
            <head>
                <style>
                    body { font-family: 'Georgia', serif; line-height: 1.7; color: #2c3e50; background: #f8f9fa; margin: 0; padding: 20px; }
                    .email-container { max-width: 650px; margin: 0 auto; background: white; border-radius: 10px; overflow: hidden; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }
                    .header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; text-align: center; }
                    .header h1 { margin: 0; font-size: 28px; font-weight: 300; }
                    .body { padding: 40px; }
                    .highlight { background: #e8f4f8; padding: 15px; border-left: 4px solid #007bff; margin: 20px 0; }
                    .button { background: #28a745; color: white; padding: 15px 30px; text-decoration: none; border-radius: 25px; display: inline-block; margin: 25px 0; font-weight: bold; }
                    .footer { background: #f8f9fa; padding: 25px; text-align: center; border-top: 1px solid #eee; }
                </style>
            </head>
            <body>
                <div class="email-container">
                    <div class="header">
                        <h1>Personal Message for {{ firstname }}</h1>
                    </div>
                    <div class="body">
                        <p>Dear {{ firstname }},</p>
                        
                        <p>I hope you don't mind me reaching out directly. I've been researching companies in the {{ industry }} sector and {{ company }} caught my attention.</p>
                        
                        <div class="highlight">
                            <strong>Here's why I'm reaching out:</strong> We've developed a unique approach that helps businesses like yours achieve remarkable results in a short time frame.
                        </div>
                        
                        <p>Unlike generic solutions, we focus on practical, actionable strategies that deliver real, measurable improvements. Our clients typically see significant changes within the first 30 days.</p>
                        
                        <p>I'd be happy to share more details about how this works and discuss your specific situation.</p>
                        
                        <center>
                            <div class="button">📞 Let's Talk: {{ phone_display }} (Please save this number for future contact)</div>
                        </center>
                        
                        <p>Could we schedule a brief conversation this week? I'm flexible with timing and can work around your schedule.</p>
                        
                        <p>Thank you for your time, {{ firstname }}.</p>
                    </div>
                    <div class="footer">
                        <p><strong>Best regards,</strong><br>Your Name</p>
                        {% if phone_display %}
                        <p>Direct: {{ phone_display }}</p>
                        {% endif %}
                    </div>
                </div>
            </body>
            </html>
            """
        ]
        
        # Select random template
        template_html = random.choice(templates)
        
        # AI enhancement
        if ai_enhance and self.gemini_model:
            try:
                prompt = f"""
                Enhance this HTML email template to make it more engaging and personalized:
                
                Lead information:
                - Name: {lead_data.get('firstname', 'there')}
                - Company: {lead_data.get('company', 'your company')}
                - Industry: {lead_data.get('industry', 'your industry')}
                
                Requirements:
                - Keep the HTML structure and CSS styling
                - Make the content more personalized and compelling
                - Add industry-specific elements when possible
                - Keep the Jinja2 template variables intact
                - Maintain professional design
                - Include engaging call-to-action
                """
                
                response = self.gemini_model.generate_content(prompt)
                enhanced_html = response.text.strip()
                
                # Basic validation
                if '<html>' in enhanced_html and '{{ firstname }}' in enhanced_html:
                    template_html = enhanced_html
                    
            except Exception as e:
                self.logger.error(f"Error enhancing HTML template with AI: {e}")
        
        # Process template
        template = Template(template_html)
        
        # Prepare phone formatting
        phone_display = ""
        phone_link = ""
        
        if phone_config.get('phone_number'):
            phone_display = format_phone_number(phone_config['phone_number'], 'display')
            phone_link = format_phone_number(phone_config['phone_number'], 'click_to_call')
        
        # Render template
        content = template.render(
            firstname=lead_data.get('firstname', 'there'),
            company=lead_data.get('company', 'your company'),
            industry=lead_data.get('industry', 'your industry'),
            phone_display=phone_display,
            phone_link=phone_link
        )
        
        return content
    
    def generate_table_format(self, lead_data: Dict[str, Any], phone_config: Dict[str, Any], 
                            ai_enhance: bool = False) -> str:
        """Generate table-formatted email content."""
        
        # Prepare service/benefit data
        services = [
            {"Service": "Business Consulting", "Benefit": "Increased Revenue", "Timeline": "30 days"},
            {"Service": "Process Optimization", "Benefit": "Reduced Costs", "Timeline": "45 days"},
            {"Service": "Marketing Strategy", "Benefit": "More Leads", "Timeline": "60 days"},
            {"Service": "Team Training", "Benefit": "Better Performance", "Timeline": "90 days"}
        ]
        
        # AI enhancement for services
        if ai_enhance and self.gemini_model:
            try:
                prompt = f"""
                Create a table of services/benefits relevant to a company in the {lead_data.get('industry', 'business')} industry.
                
                Company: {lead_data.get('company', 'the company')}
                Industry: {lead_data.get('industry', 'business')}
                
                Requirements:
                - Create 4-5 rows of relevant services
                - Each row should have: Service, Benefit, Timeline
                - Make it specific to their industry
                - Keep timelines realistic (30-90 days)
                - Format as JSON array
                """
                
                response = self.gemini_model.generate_content(prompt)
                try:
                    enhanced_services = json.loads(response.text.strip())
                    if isinstance(enhanced_services, list) and len(enhanced_services) > 0:
                        services = enhanced_services
                except:
                    pass
                    
            except Exception as e:
                self.logger.error(f"Error enhancing table data with AI: {e}")
        
        # Build HTML table
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; margin: 0; padding: 20px; }}
                .container {{ max-width: 600px; margin: 0 auto; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .header h2 {{ color: #007bff; margin-bottom: 10px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
                th {{ background-color: #f8f9fa; font-weight: bold; }}
                tr:nth-child(even) {{ background-color: #f9f9f9; }}
                .cta {{ background: #28a745; color: white; padding: 12px 24px; text-decoration: none; border-radius: 4px; display: inline-block; margin: 20px 0; }}
                .footer {{ margin-top: 30px; padding-top: 20px; border-top: 1px solid #eee; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h2>Personalized Solutions for {lead_data.get('firstname', 'You')}</h2>
                    <p>Here's how we can help {lead_data.get('company', 'your business')} succeed:</p>
                </div>
                
                <table>
                    <tr>
                        <th>Service</th>
                        <th>Benefit</th>
                        <th>Timeline</th>
                    </tr>
        """
        
        for service in services:
            html_content += f"""
                    <tr>
                        <td>{service.get('Service', '')}</td>
                        <td>{service.get('Benefit', '')}</td>
                        <td>{service.get('Timeline', '')}</td>
                    </tr>
            """
        
        # Add phone info if available
        phone_section = ""
        if phone_config.get('phone_number'):
            phone_display = format_phone_number(phone_config['phone_number'], 'display')
            phone_link = format_phone_number(phone_config['phone_number'], 'click_to_call')
            phone_section = f"""
                <p>Ready to get started? Call {phone_display} (Please save this number for future contact)</p>
            """
        
        html_content += f"""
                </table>
                
                <p>Hi {lead_data.get('firstname', 'there')},</p>
                
                <p>I've put together this customized overview of how we can help {lead_data.get('company', 'your business')} achieve better results. Each service is designed to deliver specific, measurable benefits within the timelines shown.</p>
                
                <p>Would you be interested in discussing any of these solutions in more detail? I'd be happy to answer questions and explore how we can tailor our approach to your specific needs.</p>
                
                {phone_section}
                
                <div class="footer">
                    <p>Best regards,<br>Your Name</p>
                    <p>Looking forward to helping {lead_data.get('company', 'your business')} succeed!</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def generate_html_to_image(self, lead_data: Dict[str, Any], phone_config: Dict[str, Any], 
                             ai_enhance: bool = False) -> str:
        """Generate HTML content and convert to image."""
        
        # Generate HTML content first
        html_content = self.generate_html_template(lead_data, phone_config, ai_enhance)
        
        # Create a simplified version for image conversion
        image_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ 
                    font-family: Arial, sans-serif; 
                    width: 600px; 
                    margin: 0; 
                    padding: 40px; 
                    background: white;
                    line-height: 1.6;
                }}
                .container {{ 
                    border: 2px solid #007bff; 
                    border-radius: 10px; 
                    padding: 30px; 
                    background: #f8f9fa;
                }}
                .header {{ 
                    text-align: center; 
                    margin-bottom: 20px; 
                    color: #007bff; 
                }}
                .content {{ 
                    background: white; 
                    padding: 20px; 
                    border-radius: 8px; 
                    margin: 20px 0;
                }}
                .cta {{ 
                    background: #28a745; 
                    color: white; 
                    padding: 15px 30px; 
                    text-decoration: none; 
                    border-radius: 25px; 
                    display: inline-block; 
                    margin: 20px 0; 
                    font-weight: bold;
                }}
                .footer {{ 
                    text-align: center; 
                    margin-top: 20px; 
                    color: #666; 
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>Personal Message for {lead_data.get('firstname', 'You')}</h1>
                </div>
                <div class="content">
                    <p>Dear {lead_data.get('firstname', 'there')},</p>
                    
                    <p>I hope this message finds you well. I'm reaching out because I came across {lead_data.get('company', 'your company')} and was impressed by your work in the {lead_data.get('industry', 'business')} industry.</p>
                    
                    <p>We specialize in helping businesses like yours achieve remarkable results through our proven approach. Our clients typically see significant improvements within the first 30 days.</p>
                    
                    <p>I'd love to discuss how we can help {lead_data.get('company', 'your business')} reach new heights.</p>
                    
                    <center>
                        <div class="cta">Let's Connect Today!</div>
                    </center>
                    
                    <p>Would you be available for a brief conversation this week?</p>
                </div>
                <div class="footer">
                    <p><strong>Best regards,</strong><br>Your Name</p>
        """
        
        if phone_config.get('phone_number'):
            phone_display = format_phone_number(phone_config['phone_number'], 'display')
            image_html += f"<p>📞 {phone_display}</p>"
        
        image_html += """
                </div>
            </div>
        </body>
        </html>
        """
        
        try:
            # Convert HTML to image using Selenium
            image_path = self._html_to_image_selenium(image_html)
            
            if image_path:
                # Return both HTML and image path
                return {
                    'html': html_content,
                    'image_path': image_path,
                    'content_type': 'image'
                }
            else:
                # Fallback to regular HTML
                return html_content
                
        except Exception as e:
            self.logger.error(f"Error converting HTML to image: {e}")
            return html_content
    
    def _process_spintax(self, template: str, lead_data: Dict[str, Any], phone_config: Dict[str, Any]) -> str:
        """Process spintax template and personalize content."""
        
        # Process spintax variations
        def replace_spintax(match):
            options = match.group(1).split('|')
            return random.choice(options)
        
        # Replace all spintax patterns
        content = re.sub(r'\{([^}]+)\}', replace_spintax, template)
        
        # Personalization replacements
        replacements = {
            'firstname': lead_data.get('firstname', 'there'),
            'lastname': lead_data.get('lastname', ''),
            'company': lead_data.get('company', 'your company'),
            'industry': lead_data.get('industry', 'your industry'),
            'location': lead_data.get('location', 'your area')
        }
        
        for key, value in replacements.items():
            content = content.replace(f'{{{key}}}', value)
        
        # Add phone information if configured
        if phone_config.get('include_phone') and phone_config.get('phone_number'):
            phone_text = self._format_phone_for_content(phone_config)
            
            if phone_config.get('phone_placement') == 'body':
                content += f"\n\n{phone_text}"
            # If 'attachment_only', phone will be handled separately
        
        return content.strip()
    
    def _format_phone_for_content(self, phone_config: Dict[str, Any]) -> str:
        """Format phone number for content inclusion."""
        phone_number = phone_config.get('phone_number', '')
        
        if not phone_number:
            return ""
        
        phone_display = format_phone_number(phone_number, 'display')
        phone_link = format_phone_number(phone_number, 'click_to_call')
        
        return f"Call us at {phone_display} or click to call: {phone_link}"
    
    def _html_to_image_selenium(self, html_content: str) -> Optional[str]:
        """Convert HTML to image using Selenium."""
        try:
            # Setup Chrome options
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--no-sandbox')
            chrome_options.add_argument('--disable-dev-shm-usage')
            chrome_options.add_argument('--window-size=800,1200')
            
            # Create driver
            driver = webdriver.Chrome(options=chrome_options)
            
            # Create temporary HTML file
            temp_html_path = os.path.join(os.path.dirname(__file__), 'data', 'temp', 'temp_email.html')
            os.makedirs(os.path.dirname(temp_html_path), exist_ok=True)
            
            with open(temp_html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Load HTML
            driver.get(f'file://{temp_html_path}')
            
            # Wait for page to load
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )
            
            # Take screenshot
            image_path = os.path.join(os.path.dirname(__file__), 'data', 'temp', f'email_image_{datetime.now().strftime("%Y%m%d_%H%M%S")}.png')
            driver.save_screenshot(image_path)
            
            driver.quit()
            
            # Clean up temporary HTML file
            try:
                os.remove(temp_html_path)
            except:
                pass
            
            return image_path
            
        except Exception as e:
            self.logger.error(f"Error converting HTML to image with Selenium: {e}")
            return None
    
    def create_pdf_attachment(self, lead_data: Dict[str, Any], phone_config: Dict[str, Any]) -> Optional[str]:
        """Create PDF attachment with contact information."""
        try:
            # Create PDF file path
            pdf_path = os.path.join(os.path.dirname(__file__), 'data', 'temp', f'contact_info_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf')
            os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
            
            # Create PDF
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            story.append(Paragraph("Contact Information", title_style))
            story.append(Spacer(1, 20))
            
            # Contact details
            contact_style = ParagraphStyle(
                'ContactStyle',
                parent=styles['Normal'],
                fontSize=14,
                spaceAfter=12
            )
            
            if phone_config.get('phone_number'):
                phone_display = format_phone_number(phone_config['phone_number'], 'display')
                story.append(Paragraph(f"<b>Phone:</b> {phone_display}", contact_style))
            
            if lead_data.get('email'):
                story.append(Paragraph(f"<b>Email:</b> {lead_data['email']}", contact_style))
            
            # Add personalized message
            story.append(Spacer(1, 20))
            message = f"""
            Dear {lead_data.get('firstname', 'there')},
            
            Thank you for your interest in our services. We're excited about the opportunity to help {lead_data.get('company', 'your business')} achieve its goals.
            
            Please don't hesitate to reach out using the contact information above. We look forward to hearing from you soon!
            
            Best regards,
            Your Name
            """
            
            story.append(Paragraph(message, styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            return pdf_path
            
        except Exception as e:
            self.logger.error(f"Error creating PDF attachment: {e}")
            return None
    
    def create_docx_attachment(self, lead_data: Dict[str, Any], phone_config: Dict[str, Any]) -> Optional[str]:
        """Create DOCX attachment with contact information."""
        try:
            # Create DOCX file path
            docx_path = os.path.join(os.path.dirname(__file__), 'data', 'temp', f'contact_info_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
            os.makedirs(os.path.dirname(docx_path), exist_ok=True)
            
            # Create document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Contact Information', 0)
            title.alignment = 1  # Center alignment
            
            # Add contact details
            doc.add_paragraph()
            
            if phone_config.get('phone_number'):
                phone_display = format_phone_number(phone_config['phone_number'], 'display')
                phone_para = doc.add_paragraph()
                phone_para.add_run('Phone: ').bold = True
                phone_para.add_run(phone_display)
            
            if lead_data.get('email'):
                email_para = doc.add_paragraph()
                email_para.add_run('Email: ').bold = True
                email_para.add_run(lead_data['email'])
            
            # Add personalized message
            doc.add_paragraph()
            message = f"""Dear {lead_data.get('firstname', 'there')},

Thank you for your interest in our services. We're excited about the opportunity to help {lead_data.get('company', 'your business')} achieve its goals.

Please don't hesitate to reach out using the contact information above. We look forward to hearing from you soon!

Best regards,
Your Name"""
            
            doc.add_paragraph(message)
            
            # Save document
            doc.save(docx_path)
            
            return docx_path
            
        except Exception as e:
            self.logger.error(f"Error creating DOCX attachment: {e}")
            return None
    
    def cleanup_temp_files(self) -> None:
        """Clean up temporary files."""
        try:
            temp_dir = os.path.join(os.path.dirname(__file__), 'data', 'temp')
            if os.path.exists(temp_dir):
                for filename in os.listdir(temp_dir):
                    file_path = os.path.join(temp_dir, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                    except:
                        pass
        except Exception as e:
            self.logger.error(f"Error cleaning up temp files: {e}")

def create_content_generator() -> ContentGenerator:
    """Create and return a content generator instance."""
    return ContentGenerator()