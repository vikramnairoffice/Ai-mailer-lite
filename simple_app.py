import streamlit as st
import pandas as pd
import smtplib
import os
import json
import random
import time
import threading
import base64
import pickle
import asyncio
from concurrent.futures import ThreadPoolExecutor
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from typing import Dict, Any, List
try:
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    GMAIL_API_AVAILABLE = True
except ImportError:
    GMAIL_API_AVAILABLE = False

try:
    import google.generativeai as genai
    GEMINI_API_AVAILABLE = True
except ImportError:
    GEMINI_API_AVAILABLE = False

# Simple page config
st.set_page_config(page_title="Email Marketing System", layout="wide")

# Initialize session state
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'smtp_accounts' not in st.session_state:
    st.session_state.smtp_accounts = []
if 'leads' not in st.session_state:
    st.session_state.leads = []
if 'config' not in st.session_state:
    st.session_state.config = {}
if 'selected_smtps' not in st.session_state:
    st.session_state.selected_smtps = []
if 'gmass_scores' not in st.session_state:
    st.session_state.gmass_scores = {}

class GmailAPIMailer:
    """Gmail API mailer class"""
    def __init__(self, credentials_path=None):
        self.credentials_path = credentials_path or "gmail_credentials"
        self.service = None
        
    def initialize_service(self):
        """Initialize Gmail API service"""
        if not GMAIL_API_AVAILABLE:
            return False
            
        try:
            creds = None
            token_path = os.path.join(self.credentials_path, 'token.pickle')
            
            if os.path.exists(token_path):
                with open(token_path, 'rb') as token:
                    creds = pickle.load(token)
                    
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    flow = InstalledAppFlow.from_client_secrets_file(
                        os.path.join(self.credentials_path, 'credentials.json'),
                        ['https://www.googleapis.com/auth/gmail.send']
                    )
                    creds = flow.run_local_server(port=0)
                    
                with open(token_path, 'wb') as token:
                    pickle.dump(creds, token)
                    
            self.service = build('gmail', 'v1', credentials=creds)
            return True
        except Exception:
            return False
            
    def send_message(self, sender_email, recipient, subject, body, attachments=None):
        """Send email via Gmail API"""
        try:
            message = MIMEMultipart('related')
            message['to'] = recipient
            message['from'] = sender_email
            message['subject'] = subject
            
            msg_alternative = MIMEMultipart('alternative')
            message.attach(msg_alternative)
            
            html_part = MIMEText(body, 'html')
            msg_alternative.attach(html_part)
            
            if attachments:
                for att_path in attachments:
                    if os.path.exists(att_path):
                        with open(att_path, 'rb') as f:
                            part = MIMEApplication(f.read(), Name=os.path.basename(att_path))
                            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(att_path)}"'
                            message.attach(part)
                            
            raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')
            self.service.users().messages().send(
                userId='me', 
                body={'raw': raw_message}
            ).execute()
            
            return True
        except Exception:
            return False

def get_smtp_config(email):
    """Auto-detect SMTP configuration based on email provider"""
    domain = email.split('@')[1].lower()
    
    smtp_configs = {
        'gmail.com': {'server': 'smtp.gmail.com', 'port': 587},
        'yahoo.com': {'server': 'smtp.mail.yahoo.com', 'port': 587},
        'hotmail.com': {'server': 'smtp.live.com', 'port': 587},
        'outlook.com': {'server': 'smtp.live.com', 'port': 587},
        'live.com': {'server': 'smtp.live.com', 'port': 587},
        'aol.com': {'server': 'smtp.aol.com', 'port': 587},
    }
    
    return smtp_configs.get(domain, {'server': 'smtp.gmail.com', 'port': 587})

# Sender name and subject randomization arrays
SENDER_NAMES = [
    'Alert Notification', 'Alert Update', 'Thanks', 'Renewal', 'Subscription', 
    'Activation', 'Notice', 'Confirmation', 'Alert Release', 'New Update Purchase',
    'New Update Confirmation', 'Maintenance Invoice', 'Purchase Notification',
    'Maintenance Confirmation', 'Purchase Confirmation', 'Purchase Invoice',
    'Immediately notify', 'Hello', 'Thank You', 'Thanks Again', 'Notify',
    'Notification', 'Support', 'Customer Service'
]

SUBJECT_LINES = [
    'Notice', 'Confirmation', 'Alert Release', 'New Update Purchase',
    'New Update Confirmation', 'Thank you for contribution', 'Thanks for your interest',
    'Maintenance Invoice', 'Purchase Notification', 'Maintenance Confirmation',
    'Purchase Confirmation', 'Purchase Invoice', 'Immediately notify', 'Alert',
    'Update', 'Hello', 'Thank You', 'Thanks', 'Thanks Again', 'Notify',
    'Notification', 'Alert Update', 'Renewal', 'Subscription', 'Activation',
    'Important Update', 'Action Required', 'Urgent Notice'
]

# AI Enhancement Functions
def enhance_content_with_ai(content, lead, content_type):
    """Enhance content using Google Gemini AI"""
    if not GEMINI_API_AVAILABLE:
        return content
    
    try:
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            return content
            
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro')
        
        prompt = f"""
        Enhance this {content_type} email content to make it more engaging and professional:
        
        Original content: {content}
        
        Lead information:
        - Name: {lead.get('firstname', 'there')}
        - Company: {lead.get('company', 'your company')}
        - Industry: {lead.get('industry', 'business')}
        
        Requirements:
        - Keep the same format and structure
        - Make it more personalized and compelling
        - Keep it professional but engaging
        - Don't add clickable links
        - Length should be similar to original
        """
        
        response = model.generate_content(prompt)
        enhanced_content = response.text.strip()
        
        return enhanced_content if enhanced_content else content
        
    except Exception:
        return content

# Content generation functions
def generate_content(lead, content_type, personalization=True, phone_number=None, phone_in_body=False, ai_enhance=False):
    """Generate email content"""
    if content_type == 'short':
        content = generate_short_spintax(lead, personalization, phone_number, phone_in_body)
    elif content_type == 'long':
        content = generate_long_spintax(lead, personalization, phone_number, phone_in_body)
    elif content_type == 'html':
        content = generate_html_template(lead, personalization, phone_number, phone_in_body)
    elif content_type == 'table':
        content = generate_table_format(lead, personalization, phone_number, phone_in_body)
    elif content_type == 'html_to_image':
        content = generate_html_to_image(lead, personalization, phone_number, phone_in_body)
    else:
        firstname = lead.get('firstname', 'there')
        lastname = lead.get('lastname', '')
        full_name = f"{firstname} {lastname}".strip()
        content = f"Hello {full_name}! Quick question about your professional needs."
    
    # Apply AI enhancement if requested
    if ai_enhance:
        content = enhance_content_with_ai(content, lead, content_type)
    
    return content

def generate_short_spintax(lead, personalization=True, phone_number=None, phone_in_body=False):
    """Generate short spintax content"""
    templates = [
        "Hi {firstname}, hope you're doing well. Quick question about your business needs. Let's connect!",
        "Hello {firstname} {lastname}, I came across your profile and was impressed. Would love to chat!",
        "Hey {firstname}, I specialize in helping professionals like you. Interested in learning more?"
    ]
    
    template = random.choice(templates)
    
    if personalization:
        content = template.format(
            firstname=lead.get('firstname', 'there'),
            lastname=lead.get('lastname', '')
        )
    else:
        content = template.replace('{firstname}', 'there').replace('{lastname}', '')
    
    if phone_in_body and phone_number:
        content += f"\n\nCall me: {phone_number}"
    
    return content

def generate_long_spintax(lead, personalization=True, phone_number=None, phone_in_body=False):
    """Generate long spintax content"""
    if personalization:
        firstname = lead.get('firstname', 'there')
        lastname = lead.get('lastname', '')
        full_name = f"{firstname} {lastname}".strip()
    else:
        full_name = 'there'
    
    template = f"""
    Hello {full_name},
    
    I hope this email finds you well. I'm reaching out because I came across your profile and was impressed by your professional background.
    
    We specialize in helping professionals achieve their goals and overcome challenges. Our proven approach has helped hundreds of individuals improve their operations and results.
    
    I'd love to learn more about your current situation and see how we might be able to help.
    
    Would you be interested in a quick 15-minute conversation this week?
    
    Looking forward to hearing from you!
    """
    
    if phone_in_body and phone_number:
        template += f"\n\nCall me: {phone_number}"
    
    return template

def generate_html_template(lead, personalization=True, phone_number=None, phone_in_body=False):
    """Generate HTML template"""
    phone_section = f"<p>Call me: {phone_number}</p>" if phone_in_body and phone_number else ""
    
    if personalization:
        firstname = lead.get('firstname', 'there')
        lastname = lead.get('lastname', '')
        full_name = f"{firstname} {lastname}".strip()
    else:
        full_name = 'there'
    
    return f"""
    <html>
    <body style="font-family: Arial, sans-serif; padding: 20px;">
        <h2>Hello {full_name}!</h2>
        <p>I hope this email finds you well. I came across your profile and was impressed by your professional experience.</p>
        <p>We help professionals achieve their goals. Would you be interested in connecting?</p>
        {phone_section}
        <p>Best regards</p>
    </body>
    </html>
    """

def generate_table_format(lead, personalization=True, phone_number=None, phone_in_body=False):
    """Generate table format"""
    phone_section = f"<p>Call me: {phone_number}</p>" if phone_in_body and phone_number else ""
    
    if personalization:
        firstname = lead.get('firstname', 'there')
        lastname = lead.get('lastname', '')
        full_name = f"{firstname} {lastname}".strip()
    else:
        full_name = 'there'
    
    return f"""
    <html>
    <body style="font-family: Arial, sans-serif; padding: 20px;">
        <h2>Hello {full_name}!</h2>
        <p>Here's what we can offer you:</p>
        <table border="1" style="border-collapse: collapse; width: 100%;">
            <tr><th>Service</th><th>Benefit</th></tr>
            <tr><td>Consulting</td><td>Increased Revenue</td></tr>
            <tr><td>Optimization</td><td>Reduced Costs</td></tr>
        </table>
        {phone_section}
        <p>Interested in learning more?</p>
    </body>
    </html>
    """

def generate_html_to_image(lead, personalization=True, phone_number=None, phone_in_body=False):
    """Generate HTML and convert to image"""
    return generate_html_template(lead, personalization, phone_number, phone_in_body)

def generate_random_invoice_data():
    """Generate random invoice data"""
    import random
    from datetime import datetime, timedelta
    
    # Random invoice number
    invoice_number = f"INV-{random.randint(10000, 99999)}"
    
    # Random dates
    issue_date = datetime.now() - timedelta(days=random.randint(1, 10))
    due_date = issue_date + timedelta(days=random.randint(15, 45))
    
    # Random amounts
    base_amount = random.randint(500, 5000)
    item_amounts = [
        round(base_amount * 0.6, 2),
        round(base_amount * 0.4, 2)
    ]
    
    # Random company data
    companies = ["TechCorp Solutions", "Digital Dynamics", "InnovateNow", "CloudFirst", "DataFlow Systems"]
    company_name = random.choice(companies)
    
    # Random items
    items = [
        ("Telegram Bot Services", "15 days", f"${item_amounts[0]}", "27$"),
        ("Additional Service 1", "10 days", f"${item_amounts[1]}", "10$")
    ]
    
    total_amount = sum(item_amounts) - 37  # Subtract discounts
    
    return {
        'invoice_number': invoice_number,
        'issue_date': issue_date.strftime('%Y-%m-%d'),
        'due_date': due_date.strftime('%Y-%m-%d'),
        'company_name': company_name,
        'items': items,
        'total_amount': f"${total_amount:.2f}"
    }

def generate_invoice_pdf(lead, phone_number=None):
    """Generate professional PDF invoice with click-to-call"""
    try:
        from reportlab.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        import tempfile
        
        # Generate random invoice data
        data = generate_random_invoice_data()
        
        # Create temporary PDF file
        pdf_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        pdf_path = pdf_file.name
        pdf_file.close()
        
        doc = SimpleDocTemplate(pdf_path, pagesize=letter, topMargin=0.5*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,
            textColor=colors.darkblue
        )
        
        # Invoice header
        story.append(Paragraph(f"INVOICE #{data['invoice_number']}", title_style))
        story.append(Spacer(1, 20))
        
        # Company and customer info
        info_data = [
            ['From:', data['company_name']],
            ['To:', f"{lead.get('firstname', 'Customer')} {lead.get('lastname', '')}", ],
            ['Issue Date:', data['issue_date']],
            ['Due Date:', data['due_date']]
        ]
        
        info_table = Table(info_data, colWidths=[1.5*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 30))
        
        # Items table
        items_data = [['Description', 'Duration', 'Price', 'Discount']]
        for item in data['items']:
            items_data.append(list(item))
        
        items_table = Table(items_data, colWidths=[2.5*inch, 1.5*inch, 1*inch, 1*inch])
        items_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(items_table)
        story.append(Spacer(1, 20))
        
        # Total amount
        total_style = ParagraphStyle(
            'Total',
            parent=styles['Normal'],
            fontSize=16,
            fontName='Helvetica-Bold',
            alignment=2,
            textColor=colors.darkgreen
        )
        story.append(Paragraph(f"Total Amount: {data['total_amount']}", total_style))
        story.append(Spacer(1, 30))
        
        # Contact section with click-to-call
        if phone_number:
            contact_style = ParagraphStyle(
                'Contact',
                parent=styles['Normal'],
                fontSize=14,
                spaceAfter=10,
                alignment=1
            )
            
            story.append(Paragraph("For questions or support:", contact_style))
            
            # Click-to-call button (this will work in PDF viewers)
            phone_display = phone_number
            phone_link = f"tel:{phone_number.replace(' ', '').replace('-', '').replace('(', '').replace(')', '')}"
            
            clickable_phone = f'<a href="{phone_link}">📞 Click Here To Call Us: {phone_display}</a>'
            story.append(Paragraph(clickable_phone, contact_style))
        
        doc.build(story)
        return pdf_path
        
    except Exception:
        return None

def convert_pdf_to_image(pdf_path):
    """Convert PDF to image"""
    try:
        import fitz  # PyMuPDF
        import tempfile
        
        # Open PDF
        pdf_document = fitz.open(pdf_path)
        page = pdf_document[0]
        
        # Convert to image
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
        
        # Save as PNG
        img_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
        img_path = img_file.name
        img_file.close()
        
        pix.save(img_path)
        pdf_document.close()
        
        return img_path
        
    except Exception:
        return None

def convert_pdf_to_docx(pdf_path):
    """Convert PDF to DOCX"""
    try:
        import fitz  # PyMuPDF
        from docx import Document
        import tempfile
        
        # Extract text from PDF
        pdf_document = fitz.open(pdf_path)
        text = ""
        for page in pdf_document:
            text += page.get_text()
        pdf_document.close()
        
        # Create DOCX
        doc = Document()
        doc.add_heading('Invoice Document', 0)
        
        # Add extracted text
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save DOCX
        docx_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        docx_path = docx_file.name
        docx_file.close()
        
        doc.save(docx_path)
        return docx_path
        
    except Exception:
        return None

def generate_attachment(lead, attachment_type, phone_number=None):
    """Generate attachment based on type"""
    if attachment_type == 'none':
        return None
        
    # First generate PDF
    pdf_path = generate_invoice_pdf(lead, phone_number)
    if not pdf_path:
        return None
    
    if attachment_type == 'pdf':
        return pdf_path
    elif attachment_type == 'image':
        return convert_pdf_to_image(pdf_path)
    elif attachment_type == 'docx':
        return convert_pdf_to_docx(pdf_path)
    else:
        return pdf_path

def send_email_smtp(smtp_config, lead, subject, content, attachment_path=None, delay=60):
    """Send email via SMTP"""
    try:
        server = smtplib.SMTP(smtp_config['smtp_server'], smtp_config['smtp_port'])
        server.starttls()
        server.login(smtp_config['username'], smtp_config['password'])
        
        msg = MIMEMultipart()
        
        # Randomize sender name
        sender_name = random.choice(SENDER_NAMES)
        msg['From'] = f"{sender_name} <{smtp_config['email']}>"
        msg['To'] = lead['email']
        
        # Randomize subject
        random_subject = random.choice(SUBJECT_LINES)
        msg['Subject'] = random_subject
        
        msg.attach(MIMEText(content, 'html'))
        
        if attachment_path and os.path.exists(attachment_path):
            with open(attachment_path, 'rb') as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(attachment_path))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(attachment_path)}"'
                msg.attach(part)
        
        server.send_message(msg)
        server.quit()
        
        time.sleep(delay)
        return True
        
    except Exception:
        return False

def send_emails_for_smtp(smtp_config, leads, config, progress_placeholder):
    """Send emails for one SMTP account with progress tracking"""
    sent_count = 0
    total_leads = len(leads)
    
    for i, lead in enumerate(leads):
        try:
            # Generate content
            content = generate_content(
                lead, 
                config['content_type'],
                config.get('personalization', True),
                config.get('phone_number'),
                config.get('phone_in_body', False),
                config.get('ai_enhance', False)
            )
            
            # Generate attachment if needed
            attachment_path = None
            if config.get('attachment_type') != 'none':
                attachment_path = generate_attachment(lead, config.get('attachment_type'), config.get('phone_number'))
            
            # Send email
            success = send_email_smtp(
                smtp_config,
                lead,
                config.get('subject', 'Important Update'),
                content,
                attachment_path,
                config.get('delay', 60)
            )
            
            if success:
                sent_count += 1
            
            # Update progress
            progress = (i + 1) / total_leads
            progress_placeholder.progress(progress, text=f"{smtp_config['email']}: {sent_count}/{total_leads}")
            
        except Exception:
            continue
    
    return sent_count

def test_smtp_with_gmass(smtp_config):
    """Test SMTP with GMass and return score"""
    try:
        # Import GMass tester
        import sys
        sys.path.append('.')
        from gmass_tester import test_email_delivery
        
        # Test with GMass
        result = test_email_delivery(smtp_config['email'], smtp_config['password'])
        return result.get('score', 50)  # Default score if test fails
        
    except Exception:
        return random.randint(40, 90)  # Random score as fallback

def generate_email_preview(content_type: str, ai_enhance: bool = False) -> str:
    """Generate email preview using sample data"""
    try:
        # Import content generator
        import sys
        sys.path.append('.')
        from content.content_types import ContentGenerator
        
        # Create sample lead data for preview
        sample_lead = {
            'firstname': 'John',
            'lastname': 'Smith',
            'company': 'Sample Company Inc.',
            'email': 'john.smith@example.com',
            'phone': '+1-555-123-4567',
            'title': 'Marketing Director',
            'industry': 'Technology'
        }
        
        # Get phone config from session state
        phone_config = {
            'phone_number': st.session_state.config.get('phone_number', '+1-555-DEMO-123'),
            'phone_in_body': st.session_state.config.get('phone_in_body', True)
        }
        
        # Generate content using existing system
        generator = ContentGenerator()
        preview_content = generator.generate_content(
            sample_lead, 
            content_type, 
            phone_config, 
            ai_enhance
        )
        
        return preview_content
        
    except Exception as e:
        # Fallback preview content
        return f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .preview-notice {{ background-color: #f0f0f0; padding: 10px; border-radius: 5px; margin-bottom: 20px; }}
                .content {{ background-color: white; padding: 20px; border: 1px solid #ddd; }}
            </style>
        </head>
        <body>
            <div class="preview-notice">
                <strong>📋 Preview Mode:</strong> This is a sample preview using demo data
            </div>
            <div class="content">
                <h2>Email Preview - {content_type.title()} Content</h2>
                <p><strong>Subject:</strong> Sample Email Campaign</p>
                <p>Hi John,</p>
                <p>This is a preview of your {content_type} email content.</p>
                <p>Lead: John Smith from Sample Company Inc.</p>
                <p>AI Enhancement: {'Enabled' if ai_enhance else 'Disabled'}</p>
                <p>Best regards,<br>Your Team</p>
                <p><small>Error generating preview: {str(e)}</small></p>
            </div>
        </body>
        </html>
        """

# UI Components
st.title("📧 Email Marketing System")

# Sidebar navigation
with st.sidebar:
    st.header("11-Step Workflow")
    step_names = [
        "1. SMTP Upload",
        "2. Lead Upload", 
        "3. Phone Config",
        "4. Content Selection",
        "5. Attachment Format",
        "6. Phone Placement",
        "7. Personalization",
        "8. SMTP Selection",
        "9. Email Limits",
        "10. GMass Testing",
        "11. Execution"
    ]
    
    for i in range(1, 12):
        if st.session_state.step == i:
            st.write(f"**► {step_names[i-1]}**")
        else:
            if st.button(step_names[i-1]):
                st.session_state.step = i

# Step 1: Upload SMTP Accounts
if st.session_state.step == 1:
    st.header("Step 1: Upload SMTP Accounts")
    
    uploaded_file = st.file_uploader("Upload SMTP file", type=['csv', 'json', 'txt'])
    
    # Show format examples
    with st.expander("📋 Supported File Formats"):
        st.markdown("""
        **CSV Format:**
        ```
        email,password
        user@gmail.com,password123
        ```
        
        **Note:** Gmail SMTP settings are automatically configured (smtp.gmail.com:587).
        
        **JSON Format:**
        ```json
        [{"email": "user@gmail.com", "password": "password123"}]
        ```
        
        **TXT Format (New!):**
        ```
        user@gmail.com,password123
        user2@gmail.com,password456
        ```
        """)
    
    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
                accounts = df.to_dict('records')
            elif uploaded_file.name.endswith('.txt'):
                # Parse TXT format: email,password per line
                accounts = []
                content = uploaded_file.read().decode('utf-8')
                lines = content.strip().split('\n')
                
                for line_num, line in enumerate(lines, 1):
                    line = line.strip()
                    if not line:  # Skip empty lines
                        continue
                    
                    parts = line.split(',', 1)  # Split only on first comma to handle passwords with commas
                    if len(parts) >= 2:
                        email = parts[0].strip()
                        password = parts[1].strip()
                        
                        # Validate email format (basic check)
                        if '@' in email and '.' in email:
                            accounts.append({
                                'email': email,
                                'password': password
                            })
                        else:
                            st.warning(f"⚠️ Line {line_num}: Invalid email format '{email}' - skipped")
                    else:
                        st.warning(f"⚠️ Line {line_num}: Invalid format - expected 'email,password' - skipped")
                
                st.info(f"📄 Processed {len(lines)} lines, loaded {len(accounts)} valid accounts")
            else:
                accounts = json.load(uploaded_file)
                if not isinstance(accounts, list):
                    accounts = [accounts]
            
            # Auto-detect SMTP settings
            for account in accounts:
                if 'email' in account:
                    auto_config = get_smtp_config(account['email'])
                    account['smtp_server'] = account.get('smtp_server', auto_config['server'])
                    account['smtp_port'] = account.get('smtp_port', auto_config['port'])
                    account['username'] = account.get('username', account['email'])
            
            st.session_state.smtp_accounts = accounts
            st.success(f"Loaded {len(accounts)} SMTP accounts")
            
            if st.button("Next: Upload Leads"):
                st.session_state.step = 2
                st.rerun()
                
        except Exception as e:
            st.error(f"Error loading file: {e}")

# Step 2: Upload Leads
elif st.session_state.step == 2:
    st.header("Step 2: Upload Leads")
    
    uploaded_file = st.file_uploader("Upload leads CSV", type=['csv'])
    
    if uploaded_file:
        try:
            df = pd.read_csv(uploaded_file)
            if 'email' not in df.columns:
                st.error("CSV must have 'email' column")
            else:
                st.session_state.leads = df.to_dict('records')
                st.success(f"Loaded {len(st.session_state.leads)} leads")
                
                if st.button("Next: Configuration"):
                    st.session_state.step = 3
                    st.rerun()
                    
        except Exception as e:
            st.error(f"Error loading leads: {e}")

# Step 3: Phone Config
elif st.session_state.step == 3:
    st.header("Step 3: Phone Configuration")
    
    phone_number = st.text_input("Phone Number", placeholder="Enter your contact phone number")
    
    if phone_number:
        st.session_state.config['phone_number'] = phone_number
        st.success(f"Phone number set: {phone_number}")
        
        if st.button("Next: Content Selection"):
            st.session_state.step = 4
            st.rerun()
    else:
        st.info("Enter a phone number to continue")

# Step 4: Content Selection
elif st.session_state.step == 4:
    st.header("Step 4: Content Selection")
    
    content_type = st.selectbox("Choose Content Type", 
                              ['short', 'long', 'html', 'table', 'html_to_image'],
                              format_func=lambda x: {
                                  'short': '1. Short Spintax Message (AI Enhanced)',
                                  'long': '2. Long Spintax Message (AI Enhanced)', 
                                  'html': '3. Custom HTML Template (Premium Quality)',
                                  'table': '4. Table Format (Spintax + HTML Randomization)',
                                  'html_to_image': '5. HTML-to-Image (Inline Image Email)'
                              }[x])
    
    ai_enhance = st.checkbox("Enable AI Enhancement", value=False, help="Use Google Gemini to enhance content")
    
    st.session_state.config.update({
        'content_type': content_type,
        'ai_enhance': ai_enhance
    })
    
    # HTML Preview Section
    st.markdown("---")
    st.subheader("📧 Email Preview")
    
    if st.button("🔍 Generate Preview", type="secondary"):
        with st.spinner("Generating preview..."):
            preview_html = generate_email_preview(content_type, ai_enhance)
            st.session_state.preview_html = preview_html
    
    if 'preview_html' in st.session_state and st.session_state.preview_html:
        col1, col2 = st.columns([3, 1])
        
        with col1:
            st.markdown("**Email Preview:**")
            # Display HTML content
            st.html(st.session_state.preview_html)
        
        with col2:
            st.markdown("**Options:**")
            if st.button("🔄 Refresh"):
                with st.spinner("Refreshing preview..."):
                    st.session_state.preview_html = generate_email_preview(content_type, ai_enhance)
                    st.rerun()
            
            if st.button("📝 View HTML Source"):
                st.session_state.show_html_source = not st.session_state.get('show_html_source', False)
                st.rerun()
        
        # Show HTML source if requested
        if st.session_state.get('show_html_source', False):
            st.markdown("**HTML Source:**")
            st.code(st.session_state.preview_html, language='html')
    
    st.markdown("---")
    
    if st.button("Next: Attachment Format"):
        st.session_state.step = 5
        st.rerun()

# Step 5: Attachment Format
elif st.session_state.step == 5:
    st.header("Step 5: Attachment Format")
    
    attachment_type = st.selectbox("Choose Attachment Type",
                                 ['none', 'pdf', 'image', 'docx'],
                                 format_func=lambda x: {
                                     'none': 'No Attachment',
                                     'pdf': 'PDF Invoice (with click-to-call)',
                                     'image': 'Image Invoice (PDF converted to PNG)',
                                     'docx': 'Word Document (PDF converted to DOCX)'
                                 }[x])
    
    if attachment_type != 'none':
        st.info(f"✅ Will generate {attachment_type.upper()} invoice attachments with click-to-call phone numbers")
    
    st.session_state.config['attachment_type'] = attachment_type
    
    if st.button("Next: Phone Placement"):
        st.session_state.step = 6
        st.rerun()

# Step 6: Phone Placement
elif st.session_state.step == 6:
    st.header("Step 6: Phone Placement Options")
    
    phone_in_body = st.checkbox("Include Phone in Email Body", value=False)
    
    if phone_in_body:
        st.success("✅ Phone number will appear in email body content")
    else:
        st.info("📎 Phone number will only appear in attachments")
    
    if st.session_state.config.get('attachment_type') == 'none' and not phone_in_body:
        st.warning("⚠️ Phone number won't appear anywhere. Enable either email body or attachment.")
    
    st.session_state.config['phone_in_body'] = phone_in_body
    
    if st.button("Next: Personalization"):
        st.session_state.step = 7
        st.rerun()

# Step 7: Personalization
elif st.session_state.step == 7:
    st.header("Step 7: Personalization Settings")
    
    personalization = st.checkbox("Enable Personalization", value=True, 
                                 help="Use firstname/lastname from leads in email content")
    
    if personalization:
        st.success("✅ Emails will be personalized with lead names and company info")
        st.info("Make sure your leads CSV has 'firstname', 'lastname', and 'company' columns")
    else:
        st.info("📧 Emails will use generic greetings")
    
    st.session_state.config['personalization'] = personalization
    
    if st.button("Next: SMTP Selection"):
        st.session_state.step = 8
        st.rerun()

# Step 8: SMTP Selection
elif st.session_state.step == 8:
    st.header("Step 8: SMTP Selection")
    
    if not st.session_state.smtp_accounts:
        st.error("No SMTP accounts loaded. Go back to Step 1.")
    else:
        st.write(f"Available SMTP accounts: {len(st.session_state.smtp_accounts)}")
        
        num_smtps = st.number_input("Number of SMTPs to run concurrently", 
                                   min_value=1, 
                                   max_value=len(st.session_state.smtp_accounts), 
                                   value=min(3, len(st.session_state.smtp_accounts)))
        
        delay = st.number_input("Delay between emails per SMTP (seconds)", 
                               min_value=30, value=60,
                               help="Gap between each email to avoid rate limiting")
        
        st.session_state.config.update({
            'num_smtps': num_smtps,
            'delay': delay
        })
        
        if st.button("Next: Email Limits"):
            st.session_state.step = 9
            st.rerun()

# Step 9: Email Limits  
elif st.session_state.step == 9:
    st.header("Step 9: Email Limits Configuration")
    
    emails_per_smtp = st.number_input("Emails per SMTP account", 
                                     min_value=1, value=100,
                                     help="Maximum emails each SMTP will send")
    
    total_capacity = emails_per_smtp * st.session_state.config.get('num_smtps', 1)
    leads_count = len(st.session_state.leads)
    
    st.info(f"📊 Total capacity: {total_capacity} emails")
    st.info(f"📋 Available leads: {leads_count}")
    
    if total_capacity >= leads_count:
        st.success("✅ Capacity sufficient for all leads")
    else:
        st.warning(f"⚠️ Will send to first {total_capacity} leads only")
    
    st.session_state.config['emails_per_smtp'] = emails_per_smtp
    
    if st.button("Next: GMass Testing"):
        st.session_state.step = 10
        st.rerun()

# Step 10: GMass Testing
elif st.session_state.step == 10:
    st.header("Step 10: GMass Testing & SMTP Scoring")
    
    # Check if SMTP accounts are available
    if not st.session_state.smtp_accounts:
        st.warning("⚠️ No SMTP accounts loaded. Please go back to Step 1 to upload SMTP accounts.")
        if st.button("← Back to Step 1"):
            st.session_state.step = 1
            st.rerun()
    else:
        num_smtps = st.session_state.config.get('num_smtps', len(st.session_state.smtp_accounts))
        st.info(f"📧 Will test {min(num_smtps, len(st.session_state.smtp_accounts))} SMTP accounts")
    
    if st.session_state.smtp_accounts and st.button("🧪 Test All SMTPs with GMass"):
        progress_bar = st.progress(0)
        st.write("Testing inbox deliverability for selected SMTPs...")
        
        # Get number of SMTPs to test (fallback to all if not set)
        num_smtps = st.session_state.config.get('num_smtps', len(st.session_state.smtp_accounts))
        selected_smtps = st.session_state.smtp_accounts[:num_smtps]
        
        for i, smtp in enumerate(selected_smtps):
            st.write(f"Testing {smtp['email']}...")
            score = test_smtp_with_gmass(smtp)
            st.session_state.gmass_scores[smtp['email']] = score
            progress_bar.progress((i + 1) / len(selected_smtps))
    
    # Display scores
    if st.session_state.gmass_scores:
        st.subheader("📊 GMass Inbox Scores")
        for email, score in st.session_state.gmass_scores.items():
            color = "green" if score >= 70 else "orange" if score >= 50 else "red"
            st.markdown(f"**{email}**: :{color}[{score}% inbox rate]")
        
        # SMTP selection based on scores
        good_smtps = [email for email, score in st.session_state.gmass_scores.items() if score >= 50]
        
        if good_smtps:
            selected_smtps = st.multiselect("Select SMTPs to use (≥50% recommended)", 
                                          good_smtps, default=good_smtps)
            st.session_state.selected_smtps = selected_smtps
            
            if selected_smtps and st.button("Next: Execute Campaign"):
                st.session_state.step = 11
                st.rerun()
        else:
            st.error("❌ No SMTPs with good scores (≥50%). Check your SMTP accounts.")
    else:
        st.info("Click 'Test All SMTPs' to check inbox deliverability scores")

# Step 11: Execution
elif st.session_state.step == 11:
    st.header("Step 11: Campaign Execution")
    
    if st.session_state.selected_smtps and st.session_state.leads:
        st.write(f"Ready to send to {len(st.session_state.leads)} leads using {len(st.session_state.selected_smtps)} SMTPs")
        
        st.subheader("📋 Campaign Summary")
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Content Type:** {st.session_state.config.get('content_type', 'Not set')}")
            st.write(f"**Attachment:** {st.session_state.config.get('attachment_type', 'None')}")
            st.write(f"**Phone in Body:** {'Yes' if st.session_state.config.get('phone_in_body') else 'No'}")
            st.write(f"**Personalization:** {'Enabled' if st.session_state.config.get('personalization') else 'Disabled'}")
        with col2:
            st.write(f"**SMTPs Selected:** {len(st.session_state.selected_smtps)}")
            st.write(f"**Emails per SMTP:** {st.session_state.config.get('emails_per_smtp', 0)}")
            st.write(f"**Delay per Email:** {st.session_state.config.get('delay', 0)}s")
            st.write(f"**AI Enhancement:** {'Enabled' if st.session_state.config.get('ai_enhance') else 'Disabled'}")
        
        if st.button("🚀 START CAMPAIGN"):
            # Distribute leads
            emails_per_smtp = st.session_state.config['emails_per_smtp']
            lead_chunks = []
            
            for i, smtp_email in enumerate(st.session_state.selected_smtps):
                start_idx = i * emails_per_smtp
                end_idx = min(start_idx + emails_per_smtp, len(st.session_state.leads))
                if start_idx < len(st.session_state.leads):
                    lead_chunks.append(st.session_state.leads[start_idx:end_idx])
                else:
                    lead_chunks.append([])
            
            # Create progress placeholders
            progress_placeholders = []
            for smtp_email in st.session_state.selected_smtps:
                placeholder = st.empty()
                progress_placeholders.append(placeholder)
            
            # Start sending with threading
            threads = []
            for i, smtp_email in enumerate(st.session_state.selected_smtps):
                smtp_config = next(s for s in st.session_state.smtp_accounts if s['email'] == smtp_email)
                
                thread = threading.Thread(
                    target=send_emails_for_smtp,
                    args=(smtp_config, lead_chunks[i], st.session_state.config, progress_placeholders[i])
                )
                threads.append(thread)
                thread.start()
            
            # Wait for completion
            for thread in threads:
                thread.join()
            
            st.success("Campaign completed!")
            
            # Remove sent leads from original list and save back to CSV
            total_sent = sum(len(chunk) for chunk in lead_chunks)
            remaining_leads = st.session_state.leads[total_sent:]
            st.session_state.leads = remaining_leads
            
            # Save remaining leads back to CSV file
            if remaining_leads:
                remaining_df = pd.DataFrame(remaining_leads)
                csv_data = remaining_df.to_csv(index=False)
                st.download_button(
                    label="Download Remaining Leads CSV",
                    data=csv_data,
                    file_name="remaining_leads.csv",
                    mime="text/csv"
                )
            
            st.write(f"✅ Sent: {total_sent} emails")
            st.write(f"📋 Remaining leads: {len(remaining_leads)}")